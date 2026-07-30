#!/usr/bin/env python3
"""将 hsd17b13_competition_reproducibility_proof.md 转为提交用可读 MD + DOCX。

约束：不覆盖原文件；不改 USER/ASSISTANT/TOOL 原文与哈希表数值；
仅去除 <details> 折叠并清理 AI 式外层排版。
"""

from __future__ import annotations

import argparse
import re
import sys
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parents[1]
DEFAULT_SRC = ROOT / "reports/dialogues/hsd17b13_competition_reproducibility_proof.md"
DEFAULT_MD = ROOT / "reports/dialogues/hsd17b13_competition_reproducibility_proof_submission.md"
DEFAULT_DOCX = ROOT / "reports/dialogues/hsd17b13_competition_reproducibility_proof_submission.docx"

DETAILS_RE = re.compile(
    r"<details>\n<summary><strong>(\d+) / (USER|ASSISTANT|TOOL)</strong>(.*?)</summary>\n\n(.*?)\n</details>",
    re.S,
)
SUMMARY_META_RE = re.compile(
    r"·\s*db_id=(\d+)\s*·\s*([0-9T:\.\+\-]+)\s*·\s*active=(\d+)\s*·\s*compacted=(\d+)"
)
SESSION_HDR_RE = re.compile(
    r"^### Session (\d+): `([^`]+)` — (.+)$", re.M
)
MD_TABLE_ROW_RE = re.compile(r"^\|.+\|$")
MD_TABLE_SEP_RE = re.compile(r"^\|[\s:\-|]+\|$")

COVER_SECTION_MAP = [
    ("## 披露边界", "一、披露边界"),
    ("## 证明范围", "二、证明范围"),
    ("## 会话元数据", "三、会话元数据"),
    ("## 客户端启动命令证据", "四、客户端启动命令证据"),
    ("## 工具调用统计", "五、工具调用统计"),
    ("## 关键产物哈希", "六、关键产物哈希"),
]

PART_A_LABEL_SUBS = [
    (r"\- \*\*我说\*\*：", "- 我说："),
    (r"\- \*\*Agent 判断\*\*：", "- Agent 判断："),
    (r"\- \*\*调用工具\*\*：", "- 调用工具："),
    (r"\- \*\*Agent 回应\*\*：", "- Agent 回应："),
]


@dataclass
class BodySection:
    title: str
    lang: str
    content: str


@dataclass
class AuditEntry:
    seq: str
    role: str
    db_id: str = ""
    timestamp: str = ""
    active: str = ""
    compacted: str = ""
    meta_lines: list[str] = field(default_factory=list)
    sections: list[BodySection] = field(default_factory=list)


@dataclass
class SessionBlock:
    index: int
    session_id: str
    title: str
    header_lines: list[str]
    entries: list[AuditEntry] = field(default_factory=list)


def strip_md_inline(s: str) -> str:
    s = s.replace("**", "")
    s = re.sub(r"`([^`]*)`", r"\1", s)
    return s


def parse_fenced_sections(body: str) -> tuple[list[str], list[BodySection]]:
    """解析条目 body：元数据行 + #### 标题后的围栏原文。"""
    lines = body.splitlines()
    meta: list[str] = []
    sections: list[BodySection] = []
    i = 0

    def read_fence(start: int) -> tuple[str, str, int]:
        line = lines[start]
        n = 0
        for ch in line:
            if ch == "`":
                n += 1
            else:
                break
        lang = line[n:].strip()
        start += 1
        buf: list[str] = []
        while start < len(lines):
            cur = lines[start]
            if (
                len(cur) >= n
                and cur[:n] == "`" * n
                and cur[n:].strip() == ""
                and not cur[n:].startswith("`")
            ):
                start += 1
                break
            buf.append(cur)
            start += 1
        return lang, "\n".join(buf), start

    while i < len(lines):
        line = lines[i]
        if line.startswith("#### "):
            title = line[5:].strip()
            i += 1
            while i < len(lines) and lines[i].strip() == "":
                i += 1
            if i < len(lines) and lines[i].lstrip().startswith("`"):
                lang, content, i = read_fence(i)
                sections.append(BodySection(title, lang, content))
            else:
                buf = []
                while i < len(lines) and not lines[i].startswith("#### "):
                    buf.append(lines[i])
                    i += 1
                sections.append(BodySection(title, "text", "\n".join(buf).strip("\n")))
            continue
        if line.startswith("- ") and not sections:
            meta.append(line)
            i += 1
            continue
        i += 1

    return meta, sections


def parse_summary_meta(summary_rest: str) -> dict[str, str]:
    m = SUMMARY_META_RE.search(summary_rest)
    if not m:
        return {}
    return {
        "db_id": m.group(1),
        "timestamp": m.group(2),
        "active": m.group(3),
        "compacted": m.group(4),
    }


def parse_entries(part_b: str) -> list[SessionBlock]:
    sessions: list[SessionBlock] = []
    # Split by session headers while keeping headers
    parts = SESSION_HDR_RE.split(part_b)
    # parts[0] is preamble before first session
    # then repeating: index, session_id, title, body
    if len(parts) < 5:
        raise SystemExit("Part B: 未找到 Session 分节")

    it = iter(parts[1:])
    for idx_s, sid, title, body in zip(it, it, it, it):
        # body may contain next content until end; trim trailing integrity if present
        header_lines: list[str] = []
        # extract leading meta bullets before first <details>
        before, _, after = body.partition("<details>")
        for line in before.splitlines():
            s = line.strip()
            if s.startswith("- "):
                header_lines.append(s)
        session = SessionBlock(
            index=int(idx_s),
            session_id=sid,
            title=title.strip(),
            header_lines=header_lines,
        )
        chunk = "<details>" + after if after or "<details>" in body else body
        # If partition failed because body starts differently
        if not after and "<details>" in body:
            chunk = body[body.find("<details>") :]
        # Drop anything after last details before ## 完整性
        if "## 完整性计数" in chunk:
            chunk = chunk[: chunk.find("## 完整性计数")]

        for m in DETAILS_RE.finditer(chunk if chunk.startswith("<details>") else body):
            seq, role, sum_rest, entry_body = m.group(1), m.group(2), m.group(3), m.group(4)
            sm = parse_summary_meta(sum_rest)
            meta, sections = parse_fenced_sections(entry_body)
            session.entries.append(
                AuditEntry(
                    seq=seq,
                    role=role,
                    db_id=sm.get("db_id", ""),
                    timestamp=sm.get("timestamp", ""),
                    active=sm.get("active", ""),
                    compacted=sm.get("compacted", ""),
                    meta_lines=meta,
                    sections=sections,
                )
            )
        sessions.append(session)
    return sessions


def reformat_cover(cover: str) -> str:
    out = cover
    for old, new in COVER_SECTION_MAP:
        out = out.replace(old, f"## {new}", 1)
    # blockquote note → plain paragraph（内容不变）
    out = re.sub(r"^> 注：", "注：", out, flags=re.M)
    return out.rstrip() + "\n"


def reformat_part_a(part_a: str) -> str:
    """Part A：降低标题层级、去掉标签加粗；不改正文事实。"""
    text = part_a
    # Drop nested duplicate title
    text = re.sub(
        r"^### HSD17B13 / 8G9V Scientist CLI 对话压缩记录\n\n",
        "",
        text,
        count=1,
        flags=re.M,
    )
    # blockquote meta → plain paragraphs
    text = re.sub(r"^> ", "", text, flags=re.M)
    text = text.replace("#### 最终结果", "### 最终结果")
    text = re.sub(r"^#### (\d+)\. ", r"### \1. ", text, flags=re.M)
    text = text.replace("#### 关键产物", "### 关键产物")
    text = text.replace("#### 工具说明", "### 工具说明")
    for pat, repl in PART_A_LABEL_SUBS:
        text = re.sub(pat, repl, text)
    # remove decorative horizontal rules
    text = re.sub(r"^---\s*$", "", text, flags=re.M)
    text = re.sub(r"\n{3,}", "\n\n", text)
    # rewrite Part A header
    text = text.replace(
        "## Part A：公开决策轨迹（压缩可读版）",
        "## 七、Part A：公开决策轨迹（压缩可读版）",
        1,
    )
    return text.strip() + "\n"


def format_entry_md(entry: AuditEntry) -> str:
    lines = [
        f"### 记录 {entry.seq} / {entry.role}",
        "",
    ]
    if entry.timestamp:
        lines.append(f"时间：{entry.timestamp}")
    if entry.db_id:
        lines.append(f"db_id：{entry.db_id}")
    if entry.active != "" or entry.compacted != "":
        lines.append(f"状态：active={entry.active}；compacted={entry.compacted}")
    for ml in entry.meta_lines:
        # keep original meta bullets without leading "- " duplication preference:
        # convert "- Foo：bar" -> "Foo：bar"
        lines.append(ml[2:] if ml.startswith("- ") else ml)
    lines.append("")
    for sec in entry.sections:
        lines.append(sec.title)
        lines.append("")
        fence = "```"
        # use 4 backticks if content may contain triple fences
        if "```" in sec.content:
            fence = "````"
        lang = sec.lang or "text"
        lines.append(f"{fence}{lang}")
        lines.append(sec.content)
        lines.append(fence)
        lines.append("")
    return "\n".join(lines).rstrip() + "\n"


def format_part_b_md(sessions: list[SessionBlock]) -> str:
    chunks = [
        "## 八、Part B：逐条完整审计记录",
        "",
        "以下逐条保留数据库中的原始 content、工具参数和工具返回。"
        "本提交版已去除折叠块，按会话与序号连续展开，便于 Word 阅读与人工核查。",
        "",
    ]
    for sess in sessions:
        chunks.append(
            f"### Session {sess.index}: `{sess.session_id}` — {sess.title}"
        )
        chunks.append("")
        for hl in sess.header_lines:
            chunks.append(hl)
        chunks.append("")
        for entry in sess.entries:
            chunks.append(format_entry_md(entry))
            chunks.append("")
    return "\n".join(chunks).rstrip() + "\n"


def build_submission_md(src_text: str) -> tuple[str, list[SessionBlock], str]:
    if "## Part A：" not in src_text or "## Part B：" not in src_text:
        raise SystemExit("源文件缺少 Part A / Part B 分节")

    cover = src_text.split("## Part A：", 1)[0]
    rest = src_text.split("## Part A：", 1)[1]
    part_a_body, part_b_and_tail = rest.split("## Part B：", 1)
    part_a = "## Part A：" + part_a_body
    part_b_full = "## Part B：" + part_b_and_tail

    integrity = ""
    if "## 完整性计数" in part_b_full:
        part_b_body, integrity = part_b_full.split("## 完整性计数", 1)
        integrity = "## 九、完整性计数" + integrity
        part_b_full = part_b_body

    sessions = parse_entries(part_b_full)
    title_line = cover.splitlines()[0]
    # rewrite title remains
    cover_body = "\n".join(cover.splitlines()[1:]).lstrip("\n")
    cover_fmt = reformat_cover(cover_body)
    part_a_fmt = reformat_part_a(part_a)
    part_b_fmt = format_part_b_md(sessions)

    md = "\n".join(
        [
            title_line,
            "",
            cover_fmt.strip(),
            "",
            part_a_fmt.strip(),
            "",
            part_b_fmt.strip(),
            "",
            integrity.strip(),
            "",
        ]
    )
    md = re.sub(r"\n{3,}", "\n\n", md)
    return md, sessions, cover


def set_run_font(run, *, east_asia: str, ascii_font: str, size_pt: float, bold: bool = False):
    run.bold = bold
    run.font.size = Pt(size_pt)
    run.font.name = ascii_font
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), ascii_font)
    rFonts.set(qn("w:hAnsi"), ascii_font)
    rFonts.set(qn("w:eastAsia"), east_asia)


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(6)

    for style_name, size, ea, bold in [
        ("Title", 18, "黑体", True),
        ("Heading 1", 14, "黑体", True),
        ("Heading 2", 12, "黑体", True),
        ("Heading 3", 11, "黑体", True),
    ]:
        st = styles[style_name]
        st.font.size = Pt(size)
        st.font.bold = bold
        st.font.name = "Times New Roman"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), ea)
        st.paragraph_format.space_before = Pt(12)
        st.paragraph_format.space_after = Pt(6)


def add_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header
    hp = header.paragraphs[0]
    hp.text = "HSD17B13 / 8G9V 复现证明 · 可审计记录"
    set_run_font(hp.runs[0], east_asia="宋体", ascii_font="Times New Roman", size_pt=9)
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("第 ")
    set_run_font(run, east_asia="宋体", ascii_font="Times New Roman", size_pt=9)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), " PAGE ")
    fp._p.append(fld)
    run2 = fp.add_run(" 页")
    set_run_font(run2, east_asia="宋体", ascii_font="Times New Roman", size_pt=9)


def add_para(doc: Document, text: str, *, style: str | None = None, mono: bool = False, size: float = 11):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    run = p.add_run(text)
    if mono:
        set_run_font(run, east_asia="宋体", ascii_font="Consolas", size_pt=size)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(2)
    else:
        set_run_font(run, east_asia="宋体", ascii_font="Times New Roman", size_pt=size)
    return p


def add_code_block(doc: Document, content: str) -> None:
    # Split large blocks into paragraphs per line for Word stability
    lines = content.split("\n") if content else [""]
    for line in lines:
        p = doc.add_paragraph()
        # preserve leading spaces via monospace; Word collapses less in Consolas
        run = p.add_run(line if line else " ")
        set_run_font(run, east_asia="宋体", ascii_font="Consolas", size_pt=8)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(0.3)


def parse_md_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    i = start
    while i < len(lines) and MD_TABLE_ROW_RE.match(lines[i].strip()):
        row = lines[i].strip()
        if MD_TABLE_SEP_RE.match(row):
            i += 1
            continue
        cells = [c.strip().strip("`") for c in row.strip("|").split("|")]
        rows.append(cells)
        i += 1
    return rows, i


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for ri, row in enumerate(rows):
        for ci in range(cols):
            cell = table.rows[ri].cells[ci]
            text = row[ci] if ci < len(row) else ""
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(text)
            size = 8 if len(text) > 40 else 9
            set_run_font(run, east_asia="宋体", ascii_font="Consolas" if len(text) > 20 else "Times New Roman", size_pt=size)
            if ri == 0:
                run.bold = True


def write_docx_from_structures(
    title: str,
    cover_md: str,
    part_a_md: str,
    sessions: list[SessionBlock],
    integrity_md: str,
    out_path: Path,
) -> None:
    doc = Document()
    configure_styles(doc)
    add_header_footer(doc)

    t = doc.add_paragraph(strip_md_inline(title.lstrip("# ").strip()), style="Title")
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def emit_md_chunk(md: str, *, page_break_before_h1: bool = False) -> None:
        lines = md.splitlines()
        i = 0
        first_h1 = True
        while i < len(lines):
            line = lines[i]
            if not line.strip():
                i += 1
                continue
            if line.startswith("## "):
                if page_break_before_h1 and not first_h1:
                    doc.add_page_break()
                first_h1 = False
                doc.add_heading(strip_md_inline(line[3:].strip()), level=1)
                i += 1
                continue
            if line.startswith("### "):
                doc.add_heading(strip_md_inline(line[4:].strip()), level=2)
                i += 1
                continue
            if MD_TABLE_ROW_RE.match(line.strip()):
                rows, i = parse_md_table(lines, i)
                add_table(doc, rows)
                continue
            if line.startswith("- "):
                p = doc.add_paragraph(style="List Bullet")
                run = p.add_run(strip_md_inline(line[2:]))
                set_run_font(run, east_asia="宋体", ascii_font="Times New Roman", size_pt=11)
                i += 1
                continue
            if line.startswith("```") or line.startswith("````"):
                n = 0
                for ch in line:
                    if ch == "`":
                        n += 1
                    else:
                        break
                i += 1
                buf = []
                while i < len(lines):
                    if len(lines[i]) >= n and lines[i][:n] == "`" * n and lines[i][n:].strip() == "":
                        i += 1
                        break
                    buf.append(lines[i])
                    i += 1
                add_code_block(doc, "\n".join(buf))
                continue
            add_para(doc, strip_md_inline(line))
            i += 1

    # Cover without original H1 title line
    cover_body = "\n".join(cover_md.splitlines())
    # cover_md already reformatted without # title
    emit_md_chunk(cover_body)

    emit_md_chunk(part_a_md)

    doc.add_heading("八、Part B：逐条完整审计记录", level=1)
    add_para(
        doc,
        "以下逐条保留数据库中的原始 content、工具参数和工具返回。"
        "本提交版已去除折叠块，按会话与序号连续展开，便于 Word 阅读与人工核查。",
    )

    for si, sess in enumerate(sessions):
        if si > 0:
            doc.add_page_break()
        doc.add_heading(
            f"Session {sess.index}: {sess.session_id} — {sess.title}",
            level=2,
        )
        for hl in sess.header_lines:
            add_para(doc, strip_md_inline(hl[2:] if hl.startswith("- ") else hl))

        for entry in sess.entries:
            doc.add_heading(f"记录 {entry.seq} / {entry.role}", level=3)
            meta_out = []
            if entry.timestamp:
                meta_out.append(f"时间：{entry.timestamp}")
            if entry.db_id:
                meta_out.append(f"db_id：{entry.db_id}")
            if entry.active != "" or entry.compacted != "":
                meta_out.append(f"状态：active={entry.active}；compacted={entry.compacted}")
            for ml in entry.meta_lines:
                meta_out.append(strip_md_inline(ml[2:] if ml.startswith("- ") else ml))
            for mline in meta_out:
                add_para(doc, mline, size=10)
            for sec in entry.sections:
                add_para(doc, sec.title, size=10)
                add_code_block(doc, sec.content)

    if integrity_md.strip():
        doc.add_page_break()
        emit_md_chunk(integrity_md)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def extract_hash_table(text: str) -> str:
    m = re.search(
        r"## (?:六、)?关键产物哈希\n\n(\|.+?)(?=\n## |\Z)",
        text,
        re.S,
    )
    return m.group(1).strip() if m else ""


def verify(
    src_text: str,
    md_text: str,
    sessions: list[SessionBlock],
) -> None:
    errors: list[str] = []
    if "<details>" in md_text or "</details>" in md_text or "<summary>" in md_text:
        errors.append("submission.md 仍含 details/summary")

    roles = {"USER": 0, "ASSISTANT": 0, "TOOL": 0}
    total = 0
    for sess in sessions:
        for e in sess.entries:
            roles[e.role] += 1
            total += 1
    if total != 1392:
        errors.append(f"条目总数 {total} != 1392")
    if roles != {"USER": 71, "ASSISTANT": 653, "TOOL": 668}:
        errors.append(f"角色计数异常: {roles}")

    src_hash = extract_hash_table(src_text)
    md_hash = extract_hash_table(md_text)
    # normalize heading difference
    if src_hash.replace("`", "") != md_hash.replace("`", ""):
        # compare row-wise ignoring backticks
        def rows(t: str) -> list[str]:
            return [re.sub(r"`", "", ln.strip()) for ln in t.splitlines() if ln.startswith("|")]

        if rows(src_hash) != rows(md_hash):
            errors.append("关键产物哈希表与源文件不一致")

    # sample body integrity: compare fenced contents for first USER, first TOOL, first ASSISTANT stop
    src_entries = list(DETAILS_RE.finditer(src_text))
    flat = [e for s in sessions for e in s.entries]
    if len(src_entries) != len(flat):
        errors.append(f"解析条目数 {len(flat)} != 源 details {len(src_entries)}")
    else:
        sample_idxs = [0, 1, 3, 4, 83]  # various
        # find an ASSISTANT with 可见输出
        for i, e in enumerate(flat):
            if e.role == "ASSISTANT" and any(s.title == "可见输出（原文）" for s in e.sections):
                sample_idxs.append(i)
                break
        for i in sorted(set(sample_idxs)):
            if i >= len(flat):
                continue
            src_body = src_entries[i].group(4)
            _, src_secs = parse_fenced_sections(src_body)
            dst_secs = flat[i].sections
            if len(src_secs) != len(dst_secs):
                errors.append(f"条目 {flat[i].seq} 区块数不一致 {len(src_secs)} vs {len(dst_secs)}")
                continue
            for a, b in zip(src_secs, dst_secs):
                if a.content != b.content:
                    errors.append(
                        f"条目 {flat[i].seq} 区块「{a.title}」正文与源不一致"
                    )
                if a.title != b.title:
                    errors.append(
                        f"条目 {flat[i].seq} 区块标题 {a.title!r} vs {b.title!r}"
                    )

    if errors:
        print("VERIFY FAILED:", file=sys.stderr)
        for e in errors:
            print(" -", e, file=sys.stderr)
        raise SystemExit(1)
    print(
        f"VERIFY OK: entries={total} roles={roles} details_in_md=0 "
        f"sessions={len(sessions)}"
    )


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument("--src", type=Path, default=DEFAULT_SRC)
    ap.add_argument("--md-out", type=Path, default=DEFAULT_MD)
    ap.add_argument("--docx-out", type=Path, default=DEFAULT_DOCX)
    ap.add_argument("--skip-docx", action="store_true")
    args = ap.parse_args()

    src_text = args.src.read_text(encoding="utf-8")
    md_text, sessions, cover_raw = build_submission_md(src_text)

    # Rebuild cover_fmt / part_a for docx from md sections
    args.md_out.parent.mkdir(parents=True, exist_ok=True)
    args.md_out.write_text(md_text, encoding="utf-8")
    print(f"Wrote {args.md_out} ({args.md_out.stat().st_size} bytes)")

    verify(src_text, md_text, sessions)

    if not args.skip_docx:
        # extract sections from generated md for cover/partA/integrity
        md = md_text
        # title
        title = md.splitlines()[0]
        # between title and 七、
        idx_a = md.find("## 七、Part A")
        idx_b = md.find("## 八、Part B")
        idx_i = md.find("## 九、完整性计数")
        cover_md = md[len(title) : idx_a].strip()
        part_a_md = md[idx_a:idx_b].strip()
        integrity_md = md[idx_i:].strip() if idx_i != -1 else ""

        print("Writing DOCX (large file, may take a while)...")
        write_docx_from_structures(
            title=title,
            cover_md=cover_md,
            part_a_md=part_a_md,
            sessions=sessions,
            integrity_md=integrity_md,
            out_path=args.docx_out,
        )
        print(f"Wrote {args.docx_out} ({args.docx_out.stat().st_size} bytes)")


if __name__ == "__main__":
    main()

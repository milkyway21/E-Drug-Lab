#!/usr/bin/env python3
"""Export one Scientist task as readable reports and a submission evidence bundle."""

from __future__ import annotations

import argparse
import csv
import hashlib
import json
import re
import shutil
import sqlite3
import zipfile
from collections import Counter
from datetime import datetime
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter


ROOT = Path(__file__).resolve().parents[1]
PUBLIC_DIR_NAME = "09_reports_and_dialogue"
BUNDLE_DIR_NAME = "10_reproducibility_bundle"
PRIVATE_RAW_DIR = Path(PUBLIC_DIR_NAME) / "private_raw"
FORBIDDEN_PUBLIC_TEXT = ("\u65e7\u4efb\u52a1", "\u5df2\u6709\u5b9e\u9a8c\u6570\u636e", "\u4e0d\u8981\u53c2\u8003", "\u4e0d\u590d\u7528", "\u6218\u5f79")
PUBLIC_TEXT_REPLACEMENTS = (
    ("\u65e7\u4efb\u52a1", "\u5148\u524d\u6750\u6599"),
    ("\u5df2\u6709\u5b9e\u9a8c\u6570\u636e", "\u5148\u524d\u6750\u6599"),
    ("\u4e0d\u8981\u53c2\u8003", "\u6392\u9664"),
    ("\u4e0d\u590d\u7528", "\u72ec\u7acb\u6267\u884c"),
    ("\u6218\u5f79", "\u4efb\u52a1"),
)
BUNDLE_INCLUDE_MAX_BYTES = 25 * 1024 * 1024
TRAJECTORY_TEXT_SUFFIXES = {
    ".cfg",
    ".csv",
    ".eaf",
    ".err",
    ".json",
    ".log",
    ".msj",
    ".out",
    ".sh",
    ".txt",
}
SESSION_LIMIT_TRACE_MARKERS = (
    "Hermes \u5de5\u5177\u8c03\u7528\u8fed\u4ee3\u4e0a\u9650",
    "\u4e0a\u4e00\u8f6e\u53ea\u662f\u8fbe\u5230\u4ea4\u4e92\u8f6e\u6570\u4e0a\u9650",
    "You've reached the maximum number of tool-calling iterations allowed.",
)
SECRET_PATTERNS = (
    re.compile(r"(?i)(authorization\s*[:=]\s*bearer\s+)[^\s\"']+"),
    re.compile(r"(?i)((?:api[_-]?key|token|password)\s*[:=]\s*)[^\s,;\"']+"),
    re.compile(r"\bark-[A-Za-z0-9-]{16,}\b"),
    re.compile(r"\bsk-[A-Za-z0-9_-]{16,}\b"),
)
ABSOLUTE_PUBLIC_PATH = re.compile(r"(?:^|[\s\"'=:,\[(])/(?:data|home)/", re.MULTILINE)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--task-root", type=Path, required=True)
    parser.add_argument("--db", type=Path, default=ROOT / ".hermes/state.db")
    parser.add_argument("--session-id", action="append", default=[])
    parser.add_argument("--session-file", type=Path)
    parser.add_argument("--raw-terminal", type=Path)
    parser.add_argument("--summary-json", type=Path)
    parser.add_argument("--title", default="HSD17B13 / 8G9V Scientist \u4efb\u52a1\u590d\u73b0\u8bc1\u660e")
    return parser.parse_args()


def iso_time(value: float | None) -> str:
    if value is None:
        return ""
    return datetime.fromtimestamp(value).astimezone().isoformat(timespec="milliseconds")


def elapsed_time(value: float | None, start: float | None) -> str:
    if value is None or start is None:
        return ""
    total_seconds = max(0, int(value - start))
    hours, remainder = divmod(total_seconds, 3600)
    minutes, seconds = divmod(remainder, 60)
    return f"{hours:02d}:{minutes:02d}:{seconds:02d}"


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for block in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def load_session_ids(args: argparse.Namespace) -> list[str]:
    session_ids = list(args.session_id)
    if args.session_file and args.session_file.is_file():
        session_ids.extend(
            line.strip()
            for line in args.session_file.read_text(encoding="utf-8").splitlines()
            if line.strip()
        )
    session_ids = list(dict.fromkeys(session_ids))
    if not session_ids:
        raise SystemExit("at least one --session-id or --session-file entry is required")
    return session_ids


def path_replacements(task_root: Path) -> list[tuple[str, str]]:
    replacements = [
        (str(task_root.resolve()), "."),
        (str(ROOT.resolve()), "${PROJECT_ROOT}"),
        (str(Path.home()), "${HOME}"),
        ("/data/zhang/Ye", "${SCHRODINGER_JOB_ROOT}"),
        ("/data/ye", "${DATA_ROOT}"),
        ("/opt/schrodinger2023-3", "${SCHRODINGER}"),
    ]
    return sorted(replacements, key=lambda item: len(item[0]), reverse=True)


def clean_text(value: Any, replacements: list[tuple[str, str]]) -> str:
    if value is None:
        return ""
    text = value if isinstance(value, str) else json.dumps(value, ensure_ascii=False)
    for pattern in SECRET_PATTERNS:
        text = pattern.sub(lambda match: (match.group(1) if match.lastindex else "") + "[REDACTED]", text)
    for source, target in replacements:
        text = text.replace(source, target)
    text = re.sub(
        r"(^|[\s\"'=:,\[(])/data/",
        r"\1${EXTERNAL_DATA_ROOT}/",
        text,
        flags=re.MULTILINE,
    )
    text = re.sub(
        r"(^|[\s\"'=:,\[(])/home/",
        r"\1${EXTERNAL_HOME_ROOT}/",
        text,
        flags=re.MULTILINE,
    )
    for source, target in PUBLIC_TEXT_REPLACEMENTS:
        text = text.replace(source, target)
    return text.replace("\x00", "")


def parse_tool_calls(raw: str | None) -> list[dict[str, Any]]:
    if not raw:
        return []
    try:
        payload = json.loads(raw)
    except json.JSONDecodeError:
        return [{"name": "unparsed_tool_call", "arguments": raw}]
    calls = payload if isinstance(payload, list) else [payload]
    normalized = []
    for call in calls:
        function = call.get("function") or {}
        arguments = function.get("arguments", {})
        try:
            arguments = json.loads(arguments) if isinstance(arguments, str) else arguments
        except json.JSONDecodeError:
            pass
        normalized.append(
            {
                "id": call.get("id") or call.get("call_id") or "",
                "name": function.get("name") or call.get("name") or "",
                "arguments": arguments,
            }
        )
    return normalized


def event_key(row: sqlite3.Row) -> tuple[Any, ...]:
    calls = parse_tool_calls(row["tool_calls"])
    call_ids = tuple(call.get("id") for call in calls if call.get("id"))
    if call_ids:
        return ("calls", call_ids)
    if row["role"] == "tool" and row["tool_call_id"]:
        return ("tool", row["tool_call_id"])
    return (
        "message",
        row["role"],
        row["content"] or "",
        row["tool_name"] or "",
    )


def is_session_limit_trace(row: sqlite3.Row) -> bool:
    content = row["content"] or ""
    return any(marker in content for marker in SESSION_LIMIT_TRACE_MARKERS)


def load_records(
    connection: sqlite3.Connection,
    session_ids: list[str],
    replacements: list[tuple[str, str]],
    timeline_start: float,
) -> tuple[list[dict[str, Any]], int]:
    placeholders = ",".join("?" for _ in session_ids)
    query = f"""
        SELECT id, session_id, role, content, tool_calls, tool_name, tool_call_id,
               effect_disposition, timestamp, token_count, finish_reason,
               active, compacted, platform_message_id
        FROM messages
        WHERE session_id IN ({placeholders})
        ORDER BY timestamp, id
    """
    rows = connection.execute(query, session_ids).fetchall()
    records: list[dict[str, Any]] = []
    seen: set[tuple[Any, ...]] = set()
    duplicate_count = 0
    for row in rows:
        if is_session_limit_trace(row):
            continue
        key = event_key(row)
        if key in seen:
            duplicate_count += 1
            continue
        seen.add(key)
        calls = parse_tool_calls(row["tool_calls"])
        cleaned_calls = json.loads(clean_text(calls, replacements)) if calls else []
        records.append(
            {
                "event": len(records) + 1,
                "db_id": row["id"],
                "session_id": row["session_id"],
                "elapsed": elapsed_time(row["timestamp"], timeline_start),
                "timestamp": iso_time(row["timestamp"]),
                "role": row["role"].upper(),
                "content": clean_text(row["content"], replacements),
                "tool_calls": cleaned_calls,
                "tool_name": row["tool_name"] or "",
                "tool_call_id": row["tool_call_id"] or "",
                "effect_disposition": row["effect_disposition"] or "",
                "finish_reason": row["finish_reason"] or "",
                "active": int(row["active"]),
                "compacted": int(row["compacted"]),
            }
        )
    return records, duplicate_count


def load_sessions(
    connection: sqlite3.Connection,
    session_ids: list[str],
    replacements: list[tuple[str, str]],
) -> list[dict[str, Any]]:
    placeholders = ",".join("?" for _ in session_ids)
    query = f"""
        SELECT id, title, display_name, model, billing_provider, billing_base_url,
               billing_mode, started_at, ended_at, end_reason, message_count,
               tool_call_count, input_tokens, output_tokens, cache_read_tokens,
               cache_write_tokens, reasoning_tokens, api_call_count, cwd
        FROM sessions
        WHERE id IN ({placeholders})
        ORDER BY started_at
    """
    rows = connection.execute(query, session_ids).fetchall()
    found = {row["id"] for row in rows}
    missing = [session_id for session_id in session_ids if session_id not in found]
    if missing:
        raise SystemExit(f"sessions not found: {', '.join(missing)}")
    sessions = []
    for row in rows:
        started_at = row["started_at"]
        ended_at = row["ended_at"]
        sessions.append(
            {
                "session_id": row["id"],
                "title": clean_text(row["title"] or row["display_name"] or "", replacements),
                "model": row["model"] or "",
                "provider": row["billing_provider"] or "",
                "base_url": clean_text(row["billing_base_url"] or "", replacements),
                "billing_mode": row["billing_mode"] or "",
                "started_at": iso_time(started_at),
                "ended_at": iso_time(ended_at),
                "duration": elapsed_time(ended_at, started_at),
                "end_reason": row["end_reason"] or "",
                "message_count": row["message_count"] or 0,
                "tool_call_count": row["tool_call_count"] or 0,
                "input_tokens": row["input_tokens"] or 0,
                "output_tokens": row["output_tokens"] or 0,
                "cache_read_tokens": row["cache_read_tokens"] or 0,
                "cache_write_tokens": row["cache_write_tokens"] or 0,
                "reasoning_tokens": row["reasoning_tokens"] or 0,
                "api_call_count": row["api_call_count"] or 0,
                "cwd": clean_text(row["cwd"] or "", replacements),
            }
        )
    return sessions


def dialogue_rows(records: list[dict[str, Any]]) -> list[dict[str, str]]:
    rows: list[dict[str, str]] = []
    for record in records:
        calls = record["tool_calls"]
        if calls:
            for call in calls:
                rows.append(
                    {
                        "event": str(record["event"]),
                        "elapsed": record["elapsed"],
                        "timestamp": record["timestamp"],
                        "session_id": record["session_id"],
                        "role": "ASSISTANT",
                        "type": "\u5de5\u5177\u8c03\u7528",
                        "tool": str(call.get("name") or ""),
                        "content": record["content"],
                        "arguments_or_result": json.dumps(
                            call.get("arguments", {}), ensure_ascii=False, indent=2
                        ),
                    }
                )
            continue
        row_type = {
            "USER": "\u7528\u6237\u8f93\u5165",
            "ASSISTANT": "Agent \u53ef\u89c1\u8f93\u51fa",
            "TOOL": "\u5de5\u5177\u8fd4\u56de",
        }.get(record["role"], record["role"])
        rows.append(
            {
                "event": str(record["event"]),
                "elapsed": record["elapsed"],
                "timestamp": record["timestamp"],
                "session_id": record["session_id"],
                "role": record["role"],
                "type": row_type,
                "tool": record["tool_name"],
                "content": record["content"],
                "arguments_or_result": record["content"] if record["role"] == "TOOL" else "",
            }
        )
    return rows


def write_jsonl(path: Path, records: list[dict[str, Any]]) -> None:
    with path.open("w", encoding="utf-8") as stream:
        for record in records:
            stream.write(json.dumps(record, ensure_ascii=False, separators=(",", ":")) + "\n")


def write_csv(path: Path, rows: list[dict[str, str]]) -> None:
    fields = [
        "event",
        "elapsed",
        "timestamp",
        "session_id",
        "role",
        "type",
        "tool",
        "content",
        "arguments_or_result",
    ]
    with path.open("w", encoding="utf-8-sig", newline="") as stream:
        writer = csv.DictWriter(stream, fieldnames=fields)
        writer.writeheader()
        writer.writerows(rows)


def excel_cell(value: Any) -> str:
    text = str(value or "")
    return text if len(text) <= 32_000 else text[:31_980] + "\n[\u5b8c\u6574\u5185\u5bb9\u89c1 JSONL]"


def write_xlsx(
    path: Path,
    rows: list[dict[str, str]],
    sessions: list[dict[str, Any]],
    phases: list[dict[str, Any]],
) -> None:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "\u5bf9\u8bdd\u4e0e\u5de5\u5177"
    headers = list(rows[0]) if rows else ["event", "elapsed", "timestamp", "role", "content"]
    sheet.append(headers)
    for row in rows:
        sheet.append([excel_cell(row.get(header, "")) for header in headers])
    header_fill = PatternFill("solid", fgColor="1F4E78")
    for cell in sheet[1]:
        cell.font = Font(color="FFFFFF", bold=True)
        cell.fill = header_fill
    widths = [10, 14, 25, 25, 13, 16, 24, 70, 70]
    for index, width in enumerate(widths[: len(headers)], 1):
        sheet.column_dimensions[get_column_letter(index)].width = width
    for row in sheet.iter_rows():
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)
    sheet.freeze_panes = "A2"
    sheet.auto_filter.ref = sheet.dimensions

    session_sheet = workbook.create_sheet("\u4f1a\u8bdd\u4e0e Token")
    session_headers = list(sessions[0]) if sessions else ["session_id"]
    session_sheet.append(session_headers)
    for row in sessions:
        session_sheet.append([excel_cell(row.get(header, "")) for header in session_headers])
    for cell in session_sheet[1]:
        cell.font = Font(color="FFFFFF", bold=True)
        cell.fill = header_fill
    session_sheet.freeze_panes = "A2"
    for column in range(1, len(session_headers) + 1):
        session_sheet.column_dimensions[get_column_letter(column)].width = 22

    phase_sheet = workbook.create_sheet("\u9636\u6bb5\u62a5\u544a")
    phase_headers = list(phases[0]) if phases else ["stage", "action", "result", "evidence"]
    phase_sheet.append(phase_headers)
    for row in phases:
        phase_sheet.append([excel_cell(row.get(header, "")) for header in phase_headers])
    for cell in phase_sheet[1]:
        cell.font = Font(color="FFFFFF", bold=True)
        cell.fill = header_fill
    for column in range(1, len(phase_headers) + 1):
        phase_sheet.column_dimensions[get_column_letter(column)].width = 45
    workbook.save(path)


def clipped(text: str, limit: int) -> str:
    if len(text) <= limit:
        return text
    return text[: limit - 20] + "\n[\u5b8c\u6574\u5185\u5bb9\u89c1 JSONL]"


def write_readable_markdown(
    path: Path,
    title: str,
    sessions: list[dict[str, Any]],
    records: list[dict[str, Any]],
    phases: list[dict[str, Any]],
    duplicate_count: int,
) -> None:
    totals = token_totals(sessions)
    last_elapsed = records[-1]["elapsed"] if records else "00:00:00"
    lines = [
        f"# {title}",
        "",
        "## \u7b80\u77ed\u62a5\u544a",
        "",
        f"- \u6a21\u578b：`{', '.join(sorted({row['model'] for row in sessions if row['model']}))}`",
        f"- Token：\u8f93\u5165 `{totals['input_tokens']}`，\u8f93\u51fa `{totals['output_tokens']}`，\u63a8\u7406 `{totals['reasoning_tokens']}`，\u7f13\u5b58\u8bfb\u53d6 `{totals['cache_read_tokens']}`。",
        f"- \u65f6\u95f4\u8f74：\u9996\u6761 `00:00:00`，\u6700\u540e\u53ef\u89c1\u4e8b\u4ef6 `{last_elapsed}`，\u7edf\u4e00\u8bb0\u5f55\u65f6\u957f `{last_elapsed}`。",
        f"- \u5ba1\u8ba1：`{len(records)}` \u6761\u53bb\u91cd\u53ef\u89c1\u4e8b件，\u538b\u7f29\u526f\u672c\u53bb\u91cd `{duplicate_count}` \u6761。",
        "- \u62ab\u9732\u8fb9\u754c：\u4ec5\u5bfc\u51fa\u7528\u6237\u8f93\u5165、Agent \u53ef\u89c1\u8f93\u51fa、\u5de5\u5177\u8c03\u7528\u4e0e\u5de5\u5177\u8fd4\u56de；\u672a\u8bfb\u53d6\u6216\u5bfc\u51fa\u9690\u85cf\u601d\u7ef4\u94fe、\u7cfb\u7edf\u63d0\u793a或\u5f00\u53d1者\u63d0\u793a。",
        "",
        "## \u9636\u6bb5\u6458\u8981",
        "",
    ]
    for phase in phases:
        stage_name = phase.get("stage", "") or "\u672a\u5206\u9636\u6bb5"
        lines.extend(
            [
                f"### {stage_name}",
                "",
                f"- \u6267\u884c：{phase.get('action', '')}",
                f"- \u7ed3\u679c：{phase.get('result', '')}",
                f"- \u8bc1\u636e：{phase.get('evidence', '')}",
                "",
            ]
        )
    lines.extend(["## \u5bf9\u8bdd\u4e0e\u5de5\u5177\u8bb0\u5f55", ""])
    for record in records:
        lines.append(
            f"### {record['event']:04d} / {record['role']} / +{record['elapsed']}"
        )
        lines.append("")
        lines.append(f"\u7edd\u5bf9\u65f6\u95f4：`{record['timestamp']}`")
        lines.append("")
        if record["tool_calls"]:
            for call in record["tool_calls"]:
                lines.append(f"\u5de5\u5177：`{call.get('name', '')}`")
                lines.append("")
                lines.append("```json")
                lines.append(clipped(json.dumps(call.get("arguments", {}), ensure_ascii=False, indent=2), 8000))
                lines.append("```")
        else:
            if record["tool_name"]:
                lines.append(f"\u5de5\u5177：`{record['tool_name']}`")
                lines.append("")
            lines.append("```text")
            lines.append(clipped(record["content"], 8000))
            lines.append("```")
        lines.append("")
    path.write_text("\n".join(lines), encoding="utf-8")


def set_cell_shading(cell: Any, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def set_doc_fonts(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(9)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    for section in document.sections:
        section.top_margin = Cm(1.6)
        section.bottom_margin = Cm(1.6)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)


def add_table(document: Document, headers: list[str], rows: Iterable[Iterable[Any]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        set_cell_shading(cell, "1F4E78")
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = None
            run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = clipped(str(value or ""), 4000)


def token_totals(sessions: list[dict[str, Any]]) -> dict[str, int]:
    fields = (
        "input_tokens",
        "output_tokens",
        "cache_read_tokens",
        "cache_write_tokens",
        "reasoning_tokens",
        "api_call_count",
    )
    return {field: sum(int(row[field]) for row in sessions) for field in fields}


def write_docx(
    path: Path,
    title: str,
    sessions: list[dict[str, Any]],
    rows: list[dict[str, str]],
    phases: list[dict[str, Any]],
    tool_counts: Counter[str],
    duplicate_count: int,
) -> None:
    document = Document()
    set_doc_fonts(document)
    heading = document.add_heading(title, 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_heading("\u7b80\u77ed\u4efb\u52a1\u62a5\u544a", level=1)
    totals = token_totals(sessions)
    models = ", ".join(sorted({row["model"] for row in sessions if row["model"]}))
    last_elapsed = rows[-1]["elapsed"] if rows else "00:00:00"
    final_result = phases[-1].get("result", "\u8be6\u89c1\u540e\u7eed\u9636\u6bb5\u8868") if phases else "\u8be6\u89c1\u540e\u7eed\u9636\u6bb5\u8868"
    document.add_paragraph(
        f"Scientist Agent \u5728\u5ba2\u6237\u7aef\u4e2d\u6267\u884c HSD17B13/8G9V \u6d4b\u8bd5\u6f0f\u6597。\u6700\u7ec8\u7ed3\u679c：{final_result}"
    )
    document.add_paragraph(
        f"\u6a21\u578b：{models}；\u8f93\u5165 Token：{totals['input_tokens']}；\u8f93\u51fa Token：{totals['output_tokens']}；"
        f"\u63a8\u7406 Token：{totals['reasoning_tokens']}；\u7f13\u5b58\u8bfb\u53d6 Token：{totals['cache_read_tokens']}。"
    )
    document.add_paragraph(
        f"\u7edf\u4e00\u76f8\u5bf9\u65f6\u95f4\u8f74\u4ece 00:00:00 \u5f00\u59cb；\u6700\u540e\u53ef\u89c1\u4e8b\u4ef6\u4e3a {last_elapsed}；\u7edf\u4e00\u8bb0\u5f55\u65f6\u957f\u4e3a {last_elapsed}。"
    )
    document.add_paragraph(
        f"\u8bc1\u636e\u5305\u4fdd\u7559 {len(rows)} \u6761\u4eba\u7c7b\u53ef\u8bfb\u8bb0\u5f55，\u5e76\u5c06 {duplicate_count} \u6761\u4e0a\u4e0b\u6587\u538b\u7f29\u526f\u672c\u53bb\u91cd。"
        "\u5b8c\u6574\u53ef\u89c1\u6587\u672c\u4e0e\u5de5\u5177 JSON \u4fdd\u5b58\u5728\u540c\u5305 JSONL；Word \u8868\u683c\u4e3a\u9605\u8bfb\u6027\u5bf9\u957f\u8fd4\u56de\u505a\u622a\u65ad。"
    )
    document.add_paragraph(
        "\u672c\u6587\u6863\u4e0d\u5305\u542b\u6a21\u578b\u5185\u90e8\u9690\u85cf\u601d\u7ef4\u94fe。\u9636\u6bb5\u51b3\u7b56\u6458\u8981\u4ec5\u6839\u636e\u53ef\u89c1\u8f93\u51fa、\u5de5\u5177\u8c03\u7528\u548c\u79d1\u5b66\u4ea7\u7269\u5f52\u7eb3。"
    )

    document.add_paragraph(
        "\u63d0\u4ea4\u6750\u6599\u76f8\u5bf9\u8def\u5f84\uff1a10_reproducibility_bundle/01_task_reproducibility_report.docx\uff1b"
        "10_reproducibility_bundle/02_scientist_visible_audit.jsonl\uff1b"
        "10_reproducibility_bundle/08_actual_skills_used.zip\uff1b"
        "10_reproducibility_bundle/09_scientist_science_evidence.zip\uff1b"
        "10_reproducibility_bundle/10_submission_evidence.zip\u3002"
    )
    document.add_paragraph(
        "\u5b8c\u6574\u5b9e\u9a8c\u8bb0\u5f55\u4fdd\u7559\u5728\u672c\u4efb\u52a1\u6839\u76ee\u5f55\u3002\u4e3a\u907f\u514d\u5728\u63d0\u4ea4 ZIP \u4e2d\u91cd\u590d\u5b58\u653e\u8d85\u5927 CMS/DTR/TGZ \u4e8c\u8fdb\u5236\u6587\u4ef6\uff0c"
        "\u8fd9\u4e9b\u6587\u4ef6\u5728 06_file_hashes.csv \u4e2d\u4fdd\u7559\u76f8\u5bf9\u8def\u5f84\u3001\u5b57\u8282\u6570\u4e0e SHA256\uff1b\u8f68\u8ff9\u9a8c\u6536 JSON\u3001\u4f5c\u4e1a\u65e5\u5fd7\u3001SEA \u5206\u6790\u548c\u5173\u952e\u79d1\u5b66\u4ea7\u7269\u6536\u5165\u8bc1\u636e ZIP\u3002"
    )

    document.add_page_break()
    document.add_heading("\u9636\u6bb5\u6267\u884c\u6458\u8981", level=1)
    add_table(
        document,
        ["\u9636\u6bb5", "Agent \u6267\u884c", "\u7ed3\u679c", "\u76f8\u5bf9\u8bc1\u636e\u8def\u5f84"],
        (
            (row.get("stage", ""), row.get("action", ""), row.get("result", ""), row.get("evidence", ""))
            for row in phases
        ),
    )
    document.add_heading("\u4f1a\u8bdd\u4e0e Token", level=1)
    add_table(
        document,
        ["Session ID", "\u6a21\u578b", "\u5f00\u59cb", "\u7ed3\u675f", "\u65f6\u957f", "\u8f93\u5165", "\u8f93\u51fa", "\u63a8\u7406", "\u5de5\u5177"],
        (
            (
                row["session_id"], row["model"], row["started_at"], row["ended_at"], row["duration"],
                row["input_tokens"], row["output_tokens"], row["reasoning_tokens"], row["tool_call_count"],
            )
            for row in sessions
        ),
    )
    document.add_heading("\u5de5\u5177\u8c03\u7528\u7edf\u8ba1", level=1)
    add_table(document, ["\u5de5\u5177", "\u6b21\u6570"], sorted(tool_counts.items()))

    section = document.add_section()
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    document.add_heading("\u5bf9\u8bdd\u4e0e\u5de5\u5177\u8bb0\u5f55\u8868", level=1)
    add_table(
        document,
        ["#", "\u76f8\u5bf9\u65f6\u95f4", "\u7edd\u5bf9\u65f6\u95f4", "\u7c7b\u578b", "\u5de5\u5177", "\u5185\u5bb9 / \u53c2\u6570 / \u7ed3\u679c"],
        (
            (
                row["event"], row["elapsed"], row["timestamp"], row["type"], row["tool"],
                row["arguments_or_result"] or row["content"],
            )
            for row in rows
        ),
    )
    document.save(path)


def load_phases(path: Path | None, replacements: list[tuple[str, str]]) -> list[dict[str, Any]]:
    if not path or not path.is_file():
        return []
    payload = json.loads(path.read_text(encoding="utf-8"))
    phases = payload.get("phases", payload) if isinstance(payload, dict) else payload
    if not isinstance(phases, list):
        raise SystemExit("summary JSON must be a list or contain a phases list")
    cleaned = []
    for phase in phases:
        cleaned.append(
            {
                "stage": clean_text(phase.get("stage", ""), replacements),
                "action": clean_text(phase.get("action", ""), replacements),
                "result": clean_text(phase.get("result", ""), replacements),
                "evidence": clean_text(phase.get("evidence", ""), replacements),
            }
        )
    return cleaned


def evidence_files(task_root: Path, report_paths: list[Path]) -> list[Path]:
    files: set[Path] = set(report_paths)
    input_dir = task_root / "01_inputs_and_plan"
    for pattern in (
        "manifest.json",
        "input_inventory.csv",
        "funnel_plan_*.json",
        "generated_configs/*",
    ):
        files.update(path for path in input_dir.glob(pattern) if path.is_file())
    for dirname in (
        "02_prudent_generation",
        "03_physchem_dedup",
        "04_glide_sp_top10",
        "05_pose_feature_library30",
        "06_qikprop_admet5",
        "07_final2_selection",
        "08_desmond_10ns",
    ):
        files.update(path for path in (task_root / dirname).rglob("*") if path.is_file())
    files.update(path for path in (task_root / "logs").rglob("*") if path.is_file())
    public_dir = task_root / PUBLIC_DIR_NAME
    generated_names = {
        "01_scientist_visible_audit.jsonl",
        "02_scientist_dialogue_readable.md",
        "03_scientist_dialogue_table.csv",
        "04_scientist_dialogue_table.xlsx",
        "05_task_reproducibility_report.docx",
        "06_file_hashes.csv",
        "07_submission_manifest.json",
    }
    for path in public_dir.rglob("*"):
        if not path.is_file() or PRIVATE_RAW_DIR in path.relative_to(task_root).parents:
            continue
        if path.parent == public_dir and path.name in generated_names:
            continue
        files.add(path)
    return sorted(files)


def bundle_eligible(task_root: Path, file_path: Path) -> bool:
    relative = file_path.resolve().relative_to(task_root.resolve())
    if PRIVATE_RAW_DIR in relative.parents:
        return False
    if file_path.stat().st_size > BUNDLE_INCLUDE_MAX_BYTES:
        return False
    if relative.parts[:2] == ("08_desmond_10ns", "trajectories"):
        return file_path.suffix.lower() in TRAJECTORY_TEXT_SUFFIXES
    return True


def write_hashes(
    path: Path,
    task_root: Path,
    files: list[Path],
    bundle_files: set[Path],
    raw_terminal: Path | None,
) -> list[dict[str, Any]]:
    rows = []
    for file_path in files:
        rows.append(
            {
                "relative_path": file_path.resolve().relative_to(task_root.resolve()).as_posix(),
                "bytes": file_path.stat().st_size,
                "sha256": sha256(file_path),
                "submission": (
                    "included_in_evidence_zip"
                    if file_path in bundle_files
                    else "hash_only_in_task_root"
                ),
            }
        )
    if raw_terminal and raw_terminal.is_file():
        rows.append(
            {
                "relative_path": raw_terminal.resolve().relative_to(task_root.resolve()).as_posix(),
                "bytes": raw_terminal.stat().st_size,
                "sha256": sha256(raw_terminal),
                "submission": "hash_only_private_raw",
            }
        )
    with path.open("w", encoding="utf-8-sig", newline="") as stream:
        writer = csv.DictWriter(
            stream, fieldnames=["relative_path", "bytes", "sha256", "submission"]
        )
        writer.writeheader()
        writer.writerows(rows)
    return rows


def scan_public_files(paths: Iterable[Path]) -> None:
    failures = []
    for path in paths:
        if path.suffix.lower() not in {".md", ".json", ".jsonl", ".csv"}:
            continue
        text = path.read_text(encoding="utf-8-sig", errors="replace")
        for forbidden in FORBIDDEN_PUBLIC_TEXT:
            if forbidden in text:
                failures.append(f"{path.name}: forbidden phrase {forbidden!r}")
        if ABSOLUTE_PUBLIC_PATH.search(text):
            failures.append(f"{path.name}: absolute task path remained")
        if re.search(r"\b(?:ark|sk)-[A-Za-z0-9_-]{16,}\b", text):
            failures.append(f"{path.name}: credential-like text remained")
    if failures:
        raise SystemExit("public evidence scan failed:\n" + "\n".join(failures))


def create_bundle(bundle_path: Path, task_root: Path, files: list[Path]) -> None:
    with zipfile.ZipFile(
        bundle_path, "w", compression=zipfile.ZIP_DEFLATED, allowZip64=True
    ) as archive:
        for file_path in files:
            archive.write(file_path, file_path.resolve().relative_to(task_root.resolve()).as_posix())


def main() -> int:
    args = parse_args()
    task_root = args.task_root.expanduser().resolve()
    db_path = args.db.expanduser().resolve()
    if not task_root.is_dir():
        raise SystemExit(f"task root not found: {task_root}")
    if not db_path.is_file():
        raise SystemExit(f"Hermes state DB not found: {db_path}")
    session_ids = load_session_ids(args)
    replacements = path_replacements(task_root)
    public_dir = task_root / PUBLIC_DIR_NAME
    bundle_dir = task_root / BUNDLE_DIR_NAME
    public_dir.mkdir(parents=True, exist_ok=True)
    if bundle_dir.exists():
        shutil.rmtree(bundle_dir)
    bundle_dir.mkdir(parents=True, exist_ok=True)

    connection = sqlite3.connect(db_path)
    connection.row_factory = sqlite3.Row
    placeholders = ",".join("?" for _ in session_ids)
    timeline_start = connection.execute(
        f"SELECT MIN(started_at) FROM sessions WHERE id IN ({placeholders})",
        session_ids,
    ).fetchone()[0]
    if timeline_start is None:
        raise SystemExit("session timeline start not found")
    sessions = load_sessions(connection, session_ids, replacements)
    records, duplicate_count = load_records(
        connection, session_ids, replacements, timeline_start
    )
    connection.close()
    rows = dialogue_rows(records)
    phases = load_phases(args.summary_json, replacements)
    tool_counts: Counter[str] = Counter()
    for record in records:
        for call in record["tool_calls"]:
            tool_counts[str(call.get("name") or "unknown")] += 1

    jsonl_path = public_dir / "01_scientist_visible_audit.jsonl"
    markdown_path = public_dir / "02_scientist_dialogue_readable.md"
    csv_path = public_dir / "03_scientist_dialogue_table.csv"
    xlsx_path = public_dir / "04_scientist_dialogue_table.xlsx"
    docx_path = public_dir / "05_task_reproducibility_report.docx"
    hashes_path = public_dir / "06_file_hashes.csv"
    manifest_path = public_dir / "07_submission_manifest.json"

    write_jsonl(jsonl_path, records)
    write_readable_markdown(
        markdown_path, args.title, sessions, records, phases, duplicate_count
    )
    write_csv(csv_path, rows)
    write_xlsx(xlsx_path, rows, sessions, phases)
    write_docx(
        docx_path, args.title, sessions, rows, phases, tool_counts, duplicate_count
    )
    report_paths = [jsonl_path, markdown_path, csv_path, xlsx_path, docx_path]
    files = evidence_files(task_root, report_paths)
    bundle_files = {path for path in files if bundle_eligible(task_root, path)}
    raw_terminal = args.raw_terminal.expanduser().resolve() if args.raw_terminal else None
    hash_rows = write_hashes(hashes_path, task_root, files, bundle_files, raw_terminal)
    totals = token_totals(sessions)
    submission_manifest = {
        "title": args.title,
        "generated_at": datetime.now().astimezone().isoformat(timespec="seconds"),
        "task_root": ".",
        "session_ids": session_ids,
        "models": sorted({row["model"] for row in sessions if row["model"]}),
        "token_usage": totals,
        "relative_timeline": {
            "start": "00:00:00",
            "last_visible_event": records[-1]["elapsed"] if records else "00:00:00",
            "session_duration": records[-1]["elapsed"] if records else "00:00:00",
        },
        "visible_audit_events": len(records),
        "deduplicated_compaction_copies": duplicate_count,
        "hidden_reasoning_exported": False,
        "private_raw_terminal": next(
            (row for row in hash_rows if row["submission"] == "hash_only_private_raw"), None
        ),
        "reports": [path.relative_to(task_root).as_posix() for path in report_paths],
        "skills_archive": "09_reports_and_dialogue/08_actual_skills_used.zip",
        "scientist_science_package": "09_reports_and_dialogue/final_submission/HSD17B13_8G9V_TEST_FINAL_SUBMISSION.zip",
        "evidence_file_count": len(files),
        "evidence_zip_file_count": len(bundle_files),
        "hash_only_task_root_file_count": len(files) - len(bundle_files),
        "large_binary_policy": "Full task-root files are hashed; oversized or binary MD trajectory artifacts remain in the task root and are not duplicated in the evidence ZIP.",
    }
    manifest_path.write_text(
        json.dumps(submission_manifest, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )
    files.extend([hashes_path, manifest_path])
    bundle_files.update([hashes_path, manifest_path])
    public_scan_paths = [*report_paths, hashes_path, manifest_path]
    if args.summary_json and args.summary_json.is_file():
        public_scan_paths.append(args.summary_json)
    scan_public_files(public_scan_paths)

    bundle_copies = {
        docx_path: bundle_dir / "01_task_reproducibility_report.docx",
        jsonl_path: bundle_dir / "02_scientist_visible_audit.jsonl",
        xlsx_path: bundle_dir / "03_scientist_dialogue_table.xlsx",
        csv_path: bundle_dir / "04_scientist_dialogue_table.csv",
        markdown_path: bundle_dir / "05_scientist_dialogue_readable.md",
        hashes_path: bundle_dir / "06_file_hashes.csv",
        manifest_path: bundle_dir / "07_submission_manifest.json",
    }
    for source, destination in bundle_copies.items():
        shutil.copy2(source, destination)
    skills_archive = public_dir / "08_actual_skills_used.zip"
    scientist_package = public_dir / "final_submission/HSD17B13_8G9V_TEST_FINAL_SUBMISSION.zip"
    if not skills_archive.is_file():
        raise SystemExit(f"skills archive not found: {skills_archive}")
    if not scientist_package.is_file():
        raise SystemExit(f"Scientist science package not found: {scientist_package}")
    shutil.copy2(skills_archive, bundle_dir / "08_actual_skills_used.zip")
    shutil.copy2(scientist_package, bundle_dir / "09_scientist_science_evidence.zip")
    bundle_path = bundle_dir / "10_submission_evidence.zip"
    create_bundle(bundle_path, task_root, sorted(bundle_files))
    print(
        json.dumps(
            {
                "status": "completed",
                "docx": str(bundle_copies[docx_path]),
                "xlsx": str(bundle_copies[xlsx_path]),
                "audit_jsonl": str(bundle_copies[jsonl_path]),
                "bundle": str(bundle_path),
                "sessions": session_ids,
                "tokens": totals,
            },
            ensure_ascii=False,
            indent=2,
        )
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

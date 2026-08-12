"""Incremental, evidence-linked H0-H10 report generation."""
from __future__ import annotations

import html
import json
import os
import shutil
import subprocess
import tempfile
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Iterable

from masld_agent.funnel.manifest import STAGE_ORDER, campaign_root, load_manifest, resolve_campaign_path, stage_config


REPORT_DATA_NAME = "AUTOPILOT_REPORT_DATA.json"
REPORT_MARKDOWN_NAME = "AUTOPILOT_REPORT.md"
REPORT_DOCX_NAME = "AUTOPILOT_REPORT.docx"
REPORT_PDF_NAME = "AUTOPILOT_REPORT.pdf"
IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".bmp", ".gif"}
MAX_FIGURES_PER_STAGE = 12
_CHROMIUM_UNAVAILABLE = False


def _report_root(manifest: dict[str, Any], profile: str) -> Path:
    configured = (manifest.get("stage_output_directories") or {}).get("reports")
    base = Path(configured) if configured else Path("reports")
    if not base.is_absolute():
        base = campaign_root(manifest) / base
    return base / "funnel" / profile


def _atomic_text(path: Path, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    handle, temporary_name = tempfile.mkstemp(
        dir=str(path.parent), prefix=f".{path.name}.", suffix=".tmp"
    )
    temporary = Path(temporary_name)
    try:
        with os.fdopen(handle, "w", encoding="utf-8") as stream:
            stream.write(text)
            stream.flush()
            os.fsync(stream.fileno())
        os.replace(temporary, path)
    finally:
        temporary.unlink(missing_ok=True)


def _read_data(path: Path) -> dict[str, Any]:
    if not path.is_file():
        return {"stages": {}}
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        return {"stages": {}}
    if not isinstance(data, dict):
        return {"stages": {}}
    if not isinstance(data.get("stages"), dict):
        data["stages"] = {}
    return data


def _string_paths(value: Any) -> Iterable[str]:
    if isinstance(value, str):
        yield value
    elif isinstance(value, dict):
        for item in value.values():
            yield from _string_paths(item)
    elif isinstance(value, list):
        for item in value:
            yield from _string_paths(item)


def _existing_evidence_paths(manifest: dict[str, Any], result: dict[str, Any]) -> list[Path]:
    root = campaign_root(manifest)
    paths: list[Path] = []
    for raw in _string_paths(result.get("validation") or {}):
        candidate = Path(raw)
        if not candidate.is_absolute():
            candidate = root / candidate
        try:
            resolved = candidate.resolve()
            if resolved.is_file() and resolved.is_relative_to(root.resolve()):
                paths.append(resolved)
        except (OSError, ValueError):
            continue
    return list(dict.fromkeys(paths))


def _report_relative_path(manifest: dict[str, Any], report_root: Path, raw: Any) -> str:
    value = str(raw or "")
    if not value:
        return value
    candidate = Path(value)
    if not candidate.is_absolute():
        return candidate.as_posix()
    try:
        return candidate.resolve().relative_to(campaign_root(manifest).resolve()).as_posix()
    except ValueError:
        try:
            return candidate.resolve().relative_to(report_root.resolve()).as_posix()
        except ValueError:
            return f"external/{candidate.name}"


def _figure_roots(manifest: dict[str, Any], stage: str, result: dict[str, Any]) -> list[Path]:
    roots: list[Path] = []
    config = stage_config(manifest, stage)
    for raw in config.get("figure_dirs") or []:
        roots.append(resolve_campaign_path(manifest, raw))
    for evidence in _existing_evidence_paths(manifest, result):
        roots.extend((evidence.parent, evidence.parent / "figures"))
    return list(dict.fromkeys(root for root in roots if root.is_dir()))


def _copy_figures(
    manifest: dict[str, Any], report_root: Path, stage: str, result: dict[str, Any]
) -> list[dict[str, str]]:
    figure_root = report_root / "figures" / stage
    figure_root.mkdir(parents=True, exist_ok=True)
    candidates: list[Path] = []
    for root in _figure_roots(manifest, stage, result):
        candidates.extend(
            path
            for path in root.rglob("*")
            if path.is_file() and path.suffix.lower() in IMAGE_SUFFIXES and path.stat().st_size <= 100 * 1024 * 1024
        )
    candidates = sorted(set(candidates), key=lambda path: (-path.stat().st_mtime, str(path)))
    output: list[dict[str, str]] = []
    for index, source in enumerate(candidates[:MAX_FIGURES_PER_STAGE], start=1):
        digest = __import__("hashlib").sha256(source.read_bytes()).hexdigest()[:10]
        destination = figure_root / f"{index:02d}_{digest}_{source.name}"
        if not destination.is_file():
            shutil.copy2(source, destination)
        try:
            source_label = source.relative_to(campaign_root(manifest)).as_posix()
        except ValueError:
            source_label = f"external/{source.name}"
        output.append(
            {
                "path": destination.relative_to(report_root).as_posix(),
                "source": source_label,
                "sha256": __import__("hashlib").sha256(destination.read_bytes()).hexdigest(),
            }
        )
    return output


def _number(value: Any) -> int | float | None:
    if isinstance(value, bool):
        return None
    if isinstance(value, (int, float)):
        return value
    if isinstance(value, str):
        try:
            number = float(value)
        except ValueError:
            return None
        return int(number) if number.is_integer() else number
    return None


def _observed_count(result: dict[str, Any]) -> int | float | None:
    for source in (result, result.get("validation") or {}):
        for key in ("observed_count", "valid_count", "count", "records", "valid_unique"):
            value = _number(source.get(key)) if isinstance(source, dict) else None
            if value is not None:
                return value
    return None


def _write_count_figure(
    report_root: Path, stage: str, target_count: int, observed_count: int | float | None
) -> dict[str, str] | None:
    if observed_count is None or target_count <= 0:
        return None
    try:
        from PIL import Image, ImageDraw
    except ImportError:
        return None
    figure_root = report_root / "figures" / stage
    figure_root.mkdir(parents=True, exist_ok=True)
    path = figure_root / "00_stage_counts.png"
    width, height = 1000, 420
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    maximum = max(float(target_count), float(observed_count), 1.0)
    chart_left, chart_top, chart_right, chart_bottom = 120, 50, 920, 340
    draw.line((chart_left, chart_bottom, chart_right, chart_bottom), fill="#263238", width=2)
    draw.line((chart_left, chart_top, chart_left, chart_bottom), fill="#263238", width=2)
    bars = [("planned", target_count, "#90caf9"), ("observed", observed_count, "#66bb6a")]
    bar_width = 220
    for index, (label, value, color) in enumerate(bars):
        center = chart_left + 220 + index * 360
        top = chart_bottom - int((float(value) / maximum) * (chart_bottom - chart_top))
        draw.rectangle((center - bar_width // 2, top, center + bar_width // 2, chart_bottom), fill=color, outline="#263238")
        draw.text((center - 45, chart_bottom + 18), str(label), fill="#263238")
        draw.text((center - 35, max(chart_top, top - 24)), str(value), fill="#263238")
    draw.text((chart_left, 12), f"{stage} planned vs observed", fill="#263238")
    image.save(path)
    return {
        "path": path.relative_to(report_root).as_posix(),
        "source": "generated_from_stage_counts",
        "sha256": __import__("hashlib").sha256(path.read_bytes()).hexdigest(),
    }


def _stage_record(
    manifest: dict[str, Any],
    report_root: Path,
    stage: str,
    target_count: int,
    result: dict[str, Any],
    analysis: str | None,
) -> dict[str, Any]:
    figures = _copy_figures(manifest, report_root, stage, result)
    count_figure = _write_count_figure(report_root, stage, target_count, _observed_count(result))
    if count_figure:
        figures.append(count_figure)
    validation = result.get("validation") or {}
    factual_analysis = (analysis or "").strip()
    if not factual_analysis:
        factual_analysis = (
            f"{stage} returned status `{result.get('status') or 'unknown'}` with "
            f"observed count `{_observed_count(result) if _observed_count(result) is not None else 'not reported'}`. "
            f"Artifact validation is `{(result.get('validation') or {}).get('valid')}`; "
            "no additional scientific interpretation was supplied by the monitoring agent."
        )
    return {
        "stage": stage,
        "target_count": target_count,
        "status": result.get("status"),
        "backend": result.get("backend"),
        "reused_existing": bool(result.get("reused_existing")),
        "observed_count": _observed_count(result),
        "validation_valid": validation.get("valid"),
        "analysis": factual_analysis,
        "warnings": result.get("warnings") or [],
        "error": result.get("error"),
        "evidence": [
            _report_relative_path(manifest, report_root, item.get("path"))
            for item in validation.get("evidence") or []
            if isinstance(item, dict) and item.get("path")
        ],
        "figures": figures,
        "updated_at": datetime.now(timezone.utc).isoformat(),
        "result": json.loads(json.dumps(result, ensure_ascii=False, default=str)),
    }


def _stage_heading(record: dict[str, Any]) -> str:
    status = record.get("status") or "unknown"
    return f"## {record['stage']}  {status}"


def _render_markdown(data: dict[str, Any]) -> str:
    lines = [
        "# E-Drug Lab Funnel Report",
        "",
        f"- Task: `{data.get('task_id') or data.get('campaign_id') or 'unknown'}`",
        f"- Target: `{data.get('target_id') or 'unknown'}`",
        f"- Profile: `{data.get('profile') or 'unknown'}`",
        f"- Generated: `{data.get('updated_at')}`",
        "",
        "This report is incrementally updated from validated stage artifacts. Computational evidence is not experimental confirmation.",
        "",
        "## Timeline",
        "",
        "| Stage | Status | Planned | Observed | Validation | Backend |",
        "|---|---|---:|---:|---|---|",
    ]
    stages = data.get("stages") or {}
    for stage in STAGE_ORDER:
        record = stages.get(stage)
        if not record:
            continue
        lines.append(
            f"| `{stage}` | `{record.get('status')}` | {record.get('target_count', '')} | "
            f"{record.get('observed_count', '')} | `{record.get('validation_valid')}` | `{record.get('backend') or ''}` |"
        )
    for stage in STAGE_ORDER:
        record = stages.get(stage)
        if not record:
            continue
        lines.extend(["", _stage_heading(record), ""])
        lines.append(f"- Planned count: `{record.get('target_count')}`")
        lines.append(f"- Observed count: `{record.get('observed_count')}`")
        lines.append(f"- Validation: `{record.get('validation_valid')}`")
        lines.append(f"- Reused existing artifact: `{record.get('reused_existing')}`")
        if record.get("analysis"):
            lines.extend(["", "### Analysis", "", record["analysis"]])
        if record.get("error"):
            lines.append(f"- Error: `{record['error']}`")
        if record.get("warnings"):
            lines.extend(["", "### Warnings", ""])
            lines.extend(f"- {warning}" for warning in record["warnings"])
        if record.get("evidence"):
            lines.extend(["", "### Evidence", ""])
            lines.extend(f"- `{path}`" for path in record["evidence"])
        if record.get("figures"):
            lines.extend(["", "### Figures", ""])
            for figure in record["figures"]:
                lines.append(f"- ![{figure['path']}]({figure['path']}) source=`{figure['source']}`")
    lines.extend(["", "## Provenance", "", "All figure paths and stage evidence are relative to the report directory or campaign root as recorded in the stage JSON files.", ""])
    return "\n".join(lines)


def _write_docx(path: Path, data: dict[str, Any]) -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    document = Document()
    document.core_properties.title = "E-Drug Lab Funnel Report"
    document.core_properties.subject = "Incremental H0-H10 task report"
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9)
    title = document.add_heading("E-Drug Lab Funnel Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph(
        f"Task: {data.get('task_id') or data.get('campaign_id') or 'unknown'} | "
        f"Target: {data.get('target_id') or 'unknown'} | Profile: {data.get('profile') or 'unknown'}"
    )
    document.add_paragraph(
        "This report is incrementally updated from validated artifacts. Computational evidence is not experimental confirmation."
    )
    stages = data.get("stages") or {}
    document.add_heading("Timeline", level=1)
    table = document.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    for cell, value in zip(table.rows[0].cells, ("Stage", "Status", "Planned", "Observed", "Valid", "Backend")):
        cell.text = value
    for stage in STAGE_ORDER:
        record = stages.get(stage)
        if not record:
            continue
        cells = table.add_row().cells
        values = (
            stage,
            str(record.get("status") or ""),
            str(record.get("target_count") or ""),
            str(record.get("observed_count") or ""),
            str(record.get("validation_valid")),
            str(record.get("backend") or ""),
        )
        for cell, value in zip(cells, values):
            cell.text = value
    for stage in STAGE_ORDER:
        record = stages.get(stage)
        if not record:
            continue
        document.add_heading(f"{stage}  {record.get('status') or 'unknown'}", level=1)
        document.add_paragraph(
            f"Planned: {record.get('target_count')}; observed: {record.get('observed_count')}; "
            f"validation: {record.get('validation_valid')}; reused: {record.get('reused_existing')}"
        )
        if record.get("analysis"):
            document.add_heading("Analysis", level=2)
            document.add_paragraph(record["analysis"])
        if record.get("error"):
            document.add_paragraph(f"Error: {record['error']}")
        if record.get("warnings"):
            document.add_heading("Warnings", level=2)
            for warning in record["warnings"]:
                document.add_paragraph(str(warning), style="List Bullet")
        if record.get("evidence"):
            document.add_heading("Evidence", level=2)
            for evidence in record["evidence"]:
                document.add_paragraph(str(evidence), style="List Bullet")
        for figure in record.get("figures") or []:
            image = path.parent / figure["path"]
            if image.is_file() and image.suffix.lower() in IMAGE_SUFFIXES:
                document.add_picture(str(image), width=Inches(6.2))
                document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                document.add_paragraph(f"Figure source: {figure['source']} | SHA256: {figure['sha256']}")
    document.save(path)


def _pdf_text(value: str) -> str:
    return value.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")


def _write_minimal_pdf(path: Path, data: dict[str, Any]) -> None:
    lines = ["E-Drug Lab Funnel Report", "", f"Task: {data.get('task_id') or data.get('campaign_id') or 'unknown'}"]
    for stage in STAGE_ORDER:
        record = (data.get("stages") or {}).get(stage)
        if record:
            lines.append(f"{stage}: {record.get('status')} valid={record.get('validation_valid')}")
    content_lines = ["BT", "/F1 10 Tf", "40 760 Td"]
    for line in lines[:48]:
        content_lines.append(f"({_pdf_text(line[:180])}) Tj")
        content_lines.append("0 -15 Td")
    content_lines.append("ET")
    content = "\n".join(content_lines).encode("latin-1", errors="replace")
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        b"<< /Length " + str(len(content)).encode() + b" >>\nstream\n" + content + b"\nendstream",
    ]
    payload = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for index, obj in enumerate(objects, start=1):
        offsets.append(len(payload))
        payload.extend(f"{index} 0 obj\n".encode())
        payload.extend(obj)
        payload.extend(b"\nendobj\n")
    xref = len(payload)
    payload.extend(f"xref\n0 {len(objects) + 1}\n".encode())
    payload.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        payload.extend(f"{offset:010d} 00000 n \n".encode())
    payload.extend(f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref}\n%%EOF\n".encode())
    path.write_bytes(payload)


def _write_pillow_pdf(path: Path, data: dict[str, Any]) -> None:
    """Create a dependency-light PDF with report text and copied raster figures."""
    from PIL import Image, ImageDraw, ImageFont

    page_size = (1275, 1650)
    try:
        font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 22)
        heading_font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 30)
        title_font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 38)
    except OSError:
        font = ImageFont.load_default()
        heading_font = font
        title_font = font
    pages: list[Image.Image] = []

    def new_page() -> tuple[Image.Image, ImageDraw.ImageDraw, int]:
        page = Image.new("RGB", page_size, "white")
        return page, ImageDraw.Draw(page), 70

    page, draw, y = new_page()
    draw.text((70, y), "E-Drug Lab Funnel Report", fill="#0b3d4a", font=title_font)
    y += 70
    for line in (
        f"Task: {data.get('task_id') or data.get('campaign_id') or 'unknown'}",
        f"Target: {data.get('target_id') or 'unknown'}",
        f"Profile: {data.get('profile') or 'unknown'}",
        "Computational evidence is not experimental confirmation.",
    ):
        draw.text((70, y), line, fill="#263238", font=font)
        y += 36
    pages.append(page)
    for stage in STAGE_ORDER:
        record = (data.get("stages") or {}).get(stage)
        if not record:
            continue
        page, draw, y = new_page()
        draw.text((70, y), f"{stage}  {record.get('status') or 'unknown'}", fill="#0b3d4a", font=heading_font)
        y += 55
        facts = [
            f"Planned: {record.get('target_count')}; observed: {record.get('observed_count')}",
            f"Validation: {record.get('validation_valid')}; reused: {record.get('reused_existing')}",
            f"Backend: {record.get('backend') or 'not specified'}",
        ]
        for line in facts:
            draw.text((70, y), line, fill="#263238", font=font)
            y += 34
        analysis = str(record.get("analysis") or "No agent analysis recorded.")
        y += 12
        draw.text((70, y), "Analysis", fill="#0b3d4a", font=heading_font)
        y += 42
        for line in __import__("textwrap").wrap(analysis, width=92)[:18]:
            draw.text((70, y), line, fill="#263238", font=font)
            y += 30
        for figure in (record.get("figures") or [])[:2]:
            figure_path = path.parent / figure["path"]
            if not figure_path.is_file():
                continue
            try:
                figure_image = Image.open(figure_path).convert("RGB")
                figure_image.thumbnail((1050, 520))
                if y + figure_image.height + 70 > page_size[1]:
                    pages.append(page)
                    page, draw, y = new_page()
                page.paste(figure_image, ((page_size[0] - figure_image.width) // 2, y))
                y += figure_image.height + 25
                draw.text((70, y), f"Figure: {figure['source']}", fill="#546e7a", font=font)
                y += 35
            except (OSError, ValueError):
                continue
        pages.append(page)
    if not pages:
        page, _, _ = new_page()
        pages.append(page)
    pages[0].save(path, "PDF", resolution=150.0, save_all=True, append_images=pages[1:])


def _write_pdf(path: Path, markdown: str, data: dict[str, Any], report_root: Path) -> str:
    global _CHROMIUM_UNAVAILABLE
    chromium = None if _CHROMIUM_UNAVAILABLE else next(
        (candidate for candidate in ("/usr/bin/chromium", "/usr/bin/chromium-browser", "/snap/bin/chromium") if Path(candidate).is_file()),
        None,
    )
    if chromium:
        html_path = report_root / ".AUTOPILOT_REPORT.html"
        blocks = []
        for line in markdown.splitlines():
            if line.startswith("# "):
                blocks.append(f"<h1>{html.escape(line[2:])}</h1>")
            elif line.startswith("## "):
                blocks.append(f"<h2>{html.escape(line[3:])}</h2>")
            elif line.startswith("### "):
                blocks.append(f"<h3>{html.escape(line[4:])}</h3>")
            elif line.startswith("- "):
                blocks.append(f"<p>{html.escape(line[2:])}</p>")
            elif line.startswith("|"):
                blocks.append(f"<p>{html.escape(line)}</p>")
            elif line.startswith("!"):
                start, end = line.find("(") + 1, line.find(")")
                image_path = report_root / line[start:end]
                if image_path.is_file():
                    blocks.append(f'<img src="{image_path.as_uri()}" />')
            elif line:
                blocks.append(f"<p>{html.escape(line)}</p>")
        html_text = "<html><head><meta charset='utf-8'><style>body{font-family:Arial,sans-serif;margin:32px;color:#263238}h1{color:#0b3d4a}h2{border-bottom:1px solid #90a4ae;padding-top:16px}p{font-size:10pt;line-height:1.35}img{max-width:92%;max-height:420px;display:block;margin:12px auto}</style></head><body>" + "".join(blocks) + "</body></html>"
        _atomic_text(html_path, html_text)
        result = subprocess.run(
            [chromium, "--headless", "--no-sandbox", "--disable-gpu", "--allow-file-access-from-files", f"--print-to-pdf={path}", str(html_path)],
            capture_output=True,
            text=True,
            timeout=120,
            check=False,
        )
        html_path.unlink(missing_ok=True)
        if result.returncode == 0 and path.is_file() and path.stat().st_size > 0:
            return "chromium"
        _CHROMIUM_UNAVAILABLE = True
    try:
        _write_pillow_pdf(path, data)
        return "pillow_fallback"
    except Exception:
        _write_minimal_pdf(path, data)
        return "minimal_fallback"


def update_funnel_report(
    manifest_path: str | Path,
    *,
    stage: str,
    target_count: int = 0,
    profile: str = "full",
    result: dict[str, Any] | None = None,
    analysis: str | None = None,
) -> dict[str, Any]:
    manifest = load_manifest(manifest_path)
    normalized_stage = stage.upper()
    report_root = _report_root(manifest, profile)
    report_root.mkdir(parents=True, exist_ok=True)
    data_path = report_root / REPORT_DATA_NAME
    data = _read_data(data_path)
    data.update(
        {
            "campaign_id": manifest.get("campaign_id"),
            "target_id": manifest.get("target_id"),
            "task_id": data.get("task_id") or manifest.get("campaign_id"),
            "profile": profile,
            "updated_at": datetime.now(timezone.utc).isoformat(),
        }
    )
    data.setdefault("stages", {})
    data["stages"][normalized_stage] = _stage_record(
        manifest,
        report_root,
        normalized_stage,
        target_count,
        result or {},
        analysis,
    )
    _atomic_text(data_path, json.dumps(data, ensure_ascii=False, indent=2, default=str) + "\n")
    markdown = _render_markdown(data)
    markdown_path = report_root / REPORT_MARKDOWN_NAME
    _atomic_text(markdown_path, markdown)
    docx_path = report_root / REPORT_DOCX_NAME
    docx_status = ""
    try:
        _write_docx(docx_path, data)
        docx_status = "ok"
    except Exception as exc:  # noqa: BLE001
        docx_status = f"failed:{type(exc).__name__}: {exc}"
    pdf_path = report_root / REPORT_PDF_NAME
    try:
        pdf_status = _write_pdf(pdf_path, markdown, data, report_root)
    except Exception as exc:  # noqa: BLE001
        try:
            _write_pillow_pdf(pdf_path, data)
            pdf_status = f"pillow_fallback:{type(exc).__name__}: {exc}"
        except Exception:
            _write_minimal_pdf(pdf_path, data)
            pdf_status = f"minimal_fallback:{type(exc).__name__}: {exc}"
    return {
        "status": "ok" if docx_status == "ok" and pdf_path.is_file() else "partial",
        "markdown": str(markdown_path),
        "docx": str(docx_path) if docx_path.is_file() else None,
        "pdf": str(pdf_path) if pdf_path.is_file() else None,
        "docx_status": docx_status,
        "pdf_status": pdf_status,
        "stage_count": len(data["stages"]),
        "report_data": str(data_path),
    }
